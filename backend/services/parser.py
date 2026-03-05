"""
Vextral Document Parser Service - Enhanced Extraction
Extracts text from multiple document types with high accuracy:
- PDF: Structured text + table extraction via PyMuPDF
- DOCX: Full paragraph + table extraction via python-docx
- TXT/CSV/MD: Direct text reading
- Images: AI Vision extraction via NVIDIA NIM
"""

import io
import os
import re
import base64
import fitz  # PyMuPDF
from PIL import Image
from typing import List, Dict, Any, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()


class ParserService:
    """Enhanced document parser with support for PDF, DOCX, TXT, CSV, MD, and images"""
    
    def __init__(self):
        """Initialize the parser service"""
        # Parent chunk size for maximum context in the LLM
        self.max_parent_words = 1200
        # Child chunk size for high-precision vector search
        self.max_child_words = 300
        self.min_chunk_words = 25
        
        self.ocr_min_text_chars = int(os.getenv("PDF_OCR_MIN_TEXT_CHARS", "30"))
        self.table_min_text_chars = int(os.getenv("PDF_TABLE_MIN_TEXT_CHARS", "180"))
        self.ocr_scale = float(os.getenv("PDF_OCR_SCALE", "1.7"))
        
        # Initialize NVIDIA NIM client for Vision (images only)
        api_key = os.getenv("NVIDIA_API_KEY_KIMI", os.getenv("NVIDIA_API_KEY"))
        self.client = OpenAI(
            base_url="https://integrate.api.nvidia.com/v1",
            api_key=api_key
        )
        self.vision_model = "meta/llama-3.2-11b-vision-instruct"
    
    def chunk_text_hierarchical(self, text: str) -> List[Dict[str, str]]:
        """
        Split text into Parent and Child chunks for Auto-Merging Retrieval.
        Returns a list of dicts containing 'parent_text' and 'child_text'.
        """
        text = self._clean_text(text)
        if not text:
            return []

        parents = self.chunk_text(text, self.max_parent_words)
        
        hierarchical_chunks = []
        for parent_text in parents:
            # Overlap slightly less on children since they search the parent
            children = self.chunk_text(parent_text, self.max_child_words, overlap_ratio=0.15)
            for child_text in children:
                hierarchical_chunks.append({
                    "parent_text": parent_text,
                    "child_text": child_text
                })
        
        return hierarchical_chunks

    def chunk_text(self, text: str, max_words: int = 500, overlap_ratio: float = 0.25) -> List[str]:
        """Split text into paragraph-aware chunks with overlap for better retrieval precision."""
        text = self._clean_text(text)
        if not text:
            return []

        max_words = max(60, max_words)
        overlap_words = min(int(max_words * overlap_ratio), max_words // 4)
        segment_step = max(max_words - overlap_words, 1)

        paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
        if not paragraphs:
            paragraphs = [text]

        chunks: List[str] = []
        current_words: List[str] = []

        def flush_current_chunk():
            nonlocal current_words
            if not current_words:
                return

            if len(current_words) >= self.min_chunk_words or not chunks:
                chunks.append(" ".join(current_words))

            current_words = current_words[-overlap_words:] if overlap_words > 0 else []

        for paragraph in paragraphs:
            paragraph_words = paragraph.split()
            if not paragraph_words:
                continue

            # Split very large paragraphs by sentence first.
            if len(paragraph_words) > max_words:
                sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", paragraph) if s.strip()]
                if not sentences:
                    sentences = [
                        " ".join(paragraph_words[i:i + max_words])
                        for i in range(0, len(paragraph_words), segment_step)
                    ]

                for sentence in sentences:
                    sentence_words = sentence.split()
                    if not sentence_words:
                        continue

                    if len(sentence_words) > max_words:
                        for i in range(0, len(sentence_words), segment_step):
                            segment = sentence_words[i:i + max_words]
                            if len(segment) >= self.min_chunk_words:
                                chunks.append(" ".join(segment))
                        continue

                    if len(current_words) + len(sentence_words) > max_words and current_words:
                        flush_current_chunk()

                    current_words.extend(sentence_words)
                continue

            if len(current_words) + len(paragraph_words) > max_words and current_words:
                flush_current_chunk()

            current_words.extend(paragraph_words)

        if current_words and (len(current_words) >= self.min_chunk_words or not chunks):
            chunks.append(" ".join(current_words))

        # Remove exact duplicate chunks after normalization.
        deduped_chunks: List[str] = []
        seen: set[str] = set()
        for chunk in chunks:
            normalized = " ".join(chunk.lower().split())
            if normalized in seen:
                continue
            seen.add(normalized)
            deduped_chunks.append(chunk)

        return deduped_chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text for better quality"""
        # Remove excessive whitespace but keep paragraph breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove excessive spaces
        text = re.sub(r' {3,}', ' ', text)
        # Remove null characters
        text = text.replace('\x00', '')
        # Strip leading/trailing whitespace
        text = text.strip()
        return text
    
    def _extract_text_with_ai(self, image_base64: str) -> str:
        """Use AI Vision to transcribe image content"""
        try:
            response = self.client.chat.completions.create(
                model=self.vision_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Transcribe all text, tables, and charts from this page into clear Markdown. If there are tables, represent them as Markdown tables."},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_base64}"}}
                        ]
                    }
                ],
                max_tokens=4096,
                temperature=0.2,
                top_p=1.0
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"⚠️ AI Vision extraction failed: {e}")
            return ""

    def parse_document(self, file_bytes: bytes, filename: str) -> List[Dict]:
        """Parse any supported document into chunks"""
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'pdf':
            return self._parse_pdf(file_bytes, filename)
        elif file_ext == 'docx':
            return self._parse_docx(file_bytes, filename)
        elif file_ext in ['txt', 'csv', 'md', 'json']:
            return self._parse_text(file_bytes, filename)
        elif file_ext in ['png', 'jpg', 'jpeg', 'webp']:
            return self._parse_image(file_bytes, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: pdf, docx, txt, csv, md, json, png, jpg, jpeg, webp")
    
    def _parse_pdf(self, file_bytes: bytes, filename: str) -> List[Dict]:
        """Enhanced PDF extraction with parallel OCR and hierarchical chunking"""
        chunks = []
        chunk_index = 0
        
        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            total_pages = len(pdf_document)
            
            print(f"⚙️ Enhanced PDF Parse: {filename} ({total_pages} pages)")
            
            # Extract standard text and tables first
            pages_text = {}
            pages_needing_ocr = []
            
            for page_num in range(total_pages):
                page = pdf_document[page_num]
                page_text_parts = []
                
                # 1. Extract structured text blocks
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                
                for block in blocks:
                    if block["type"] == 0:  # Text block
                        block_text = ""
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if text:
                                    font_size = span["size"]
                                    is_bold = "bold" in span["font"].lower() or "Bold" in span["font"]
                                    
                                    if font_size > 14 and is_bold:
                                        text = f"\n## {text}\n"
                                    elif font_size > 12 and is_bold:
                                        text = f"\n### {text}\n"
                                    elif is_bold:
                                        text = f"**{text}**"
                                    
                                    line_text += text + " "
                            
                            if line_text.strip():
                                block_text += line_text.strip() + "\n"
                        
                        if block_text.strip():
                            page_text_parts.append(block_text.strip())

                base_page_text = "\n\n".join(page_text_parts)

                # 2. Extract tables only if text exists
                if len(base_page_text.strip()) >= self.table_min_text_chars:
                    try:
                        tables = page.find_tables()
                        if tables and tables.tables:
                            for table in tables.tables:
                                table_data = table.extract()
                                if table_data:
                                    md_table = self._table_to_markdown(table_data)
                                    if md_table:
                                        page_text_parts.append(f"\n{md_table}\n")
                                        # print(f"   📊 Found table on page {page_num + 1}")
                    except Exception:
                        pass

                full_page_text = "\n\n".join(page_text_parts)
                
                if not full_page_text or len(full_page_text.strip()) < self.ocr_min_text_chars:
                    if len(page.get_images(full=True)) > 0:
                        pages_needing_ocr.append((page_num, page))
                    else:
                        pages_text[page_num] = full_page_text
                else:
                    pages_text[page_num] = full_page_text

            # 3. Parallel OCR Processing for missing text using Vision API
            if pages_needing_ocr:
                print(f"   ⚠️ Running parallel OCR on {len(pages_needing_ocr)} pages...")
                
                def process_ocr_page(page_num, page):
                    try:
                        scale = min(max(self.ocr_scale, 1.2), 2.0)
                        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                        image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                        jpeg_buffer = io.BytesIO()
                        image.save(jpeg_buffer, format="JPEG", quality=75, optimize=True)
                        page_image_base64 = base64.b64encode(jpeg_buffer.getvalue()).decode("utf-8")
                        ocr_text = self._extract_text_with_ai(page_image_base64)
                        return page_num, ocr_text if ocr_text else ""
                    except Exception as e:
                        print(f"   ⚠️ OCR fallback failed on page {page_num + 1}: {e}")
                        return page_num, ""

                # Target blazing fast speed by doing max parallel workers allowed by free APIs (usually 5-10)
                max_workers = min(10, len(pages_needing_ocr))
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = {executor.submit(process_ocr_page, p_num, p): p_num for p_num, p in pages_needing_ocr}
                    for future in as_completed(futures):
                        p_num, text = future.result()
                        pages_text[p_num] = text

            # 4. Perform Hierarchical Chunking
            for page_num in sorted(pages_text.keys()):
                full_page_text = pages_text[page_num]
                if not full_page_text or len(full_page_text.strip()) < 10:
                    continue
                
                hierarchical_segments = self.chunk_text_hierarchical(full_page_text)
                
                for segment in hierarchical_segments:
                    chunks.append({
                        "text": segment["child_text"],  # Target text for embedding
                        "parent_text": segment["parent_text"],  # Full context for LLM
                        "parent_id": f"parent_{page_num}_{hash(segment['parent_text'])}",
                        "image_base64": None,
                        "page_number": page_num + 1,
                        "chunk_index": chunk_index,
                        "source_file": filename,
                        "chunk_type": "text"
                    })
                    chunk_index += 1
                
            pdf_document.close()
            print(f"✓ Parsed PDF Hierarchically: {len(chunks)} chunks from {filename}")
            return chunks
            
        except Exception as e:
            print(f"Error parsing PDF {filename}: {e}")
            raise Exception(f"Failed to parse PDF: {str(e)}")
    
    def _table_to_markdown(self, table_data: list) -> str:
        """Convert table data to Markdown table format"""
        if not table_data or len(table_data) < 2:
            return ""
        
        # Clean cells
        cleaned = []
        for row in table_data:
            cleaned.append([str(cell).strip() if cell else "" for cell in row])
        
        # Build markdown
        header = "| " + " | ".join(cleaned[0]) + " |"
        separator = "| " + " | ".join(["---"] * len(cleaned[0])) + " |"
        rows = [f"| {' | '.join(row)} |" for row in cleaned[1:]]
        
        return "\n".join([header, separator] + rows)
    
    def _parse_docx(self, file_bytes: bytes, filename: str) -> List[Dict]:
        """Parse DOCX files with paragraphs and tables"""
        try:
            from docx import Document
        except ImportError:
            raise Exception("python-docx not installed. Run: pip install python-docx")
        
        chunks = []
        chunk_index = 0
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            print(f"⚙️ Parse DOCX: {filename}")
            
            all_text_parts = []
            
            # Extract paragraphs with heading detection
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect headings
                if para.style.name.startswith('Heading 1'):
                    all_text_parts.append(f"\n## {text}\n")
                elif para.style.name.startswith('Heading 2'):
                    all_text_parts.append(f"\n### {text}\n")
                elif para.style.name.startswith('Heading'):
                    all_text_parts.append(f"\n#### {text}\n")
                else:
                    all_text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                md_table = self._table_to_markdown(table_data)
                if md_table:
                    all_text_parts.append(f"\n{md_table}\n")
            
            full_text = "\n\n".join(all_text_parts)
            hierarchical_segments = self.chunk_text_hierarchical(full_text)
            
            for segment in hierarchical_segments:
                chunks.append({
                    "text": segment["child_text"],
                    "parent_text": segment["parent_text"],
                    "parent_id": f"parent_docx_{hash(segment['parent_text'])}",
                    "image_base64": None,
                    "page_number": 1,
                    "chunk_index": chunk_index,
                    "source_file": filename,
                    "chunk_type": "text"
                })
                chunk_index += 1
            
            print(f"✓ Parsed DOCX: {len(chunks)} hierarchical chunks from {filename}")
            return chunks
            
        except Exception as e:
            print(f"Error parsing DOCX {filename}: {e}")
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    def _parse_text(self, file_bytes: bytes, filename: str) -> List[Dict]:
        """Parse plain text files (TXT, CSV, MD, JSON)"""
        chunks = []
        chunk_index = 0
        
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = file_bytes.decode('latin-1')
            
            print(f"⚙️ Parse Text: {filename}")
            
            hierarchical_segments = self.chunk_text_hierarchical(text)
            
            for segment in hierarchical_segments:
                chunks.append({
                    "text": segment["child_text"],
                    "parent_text": segment["parent_text"],
                    "parent_id": f"parent_txt_{hash(segment['parent_text'])}",
                    "image_base64": None,
                    "page_number": 1,
                    "chunk_index": chunk_index,
                    "source_file": filename,
                    "chunk_type": "text"
                })
                chunk_index += 1
            
            print(f"✓ Parsed Text: {len(chunks)} hierarchical chunks from {filename}")
            return chunks
            
        except Exception as e:
            print(f"Error parsing text {filename}: {e}")
            raise Exception(f"Failed to parse text file: {str(e)}")
    
    def _parse_image(self, file_bytes: bytes, filename: str) -> List[Dict]:
        """Parse image file using AI Vision"""
        try:
            image_base64 = base64.b64encode(file_bytes).decode('utf-8')
            
            print(f"⚙️ Parse Image with AI: {filename}")
            extracted_text = self._extract_text_with_ai(image_base64)
            
            if not extracted_text:
                extracted_text = "[Image could not be transcribed]"
            
            chunk = {
                "text": extracted_text,
                "image_base64": image_base64,
                "page_number": 1,
                "chunk_index": 0,
                "source_file": filename,
                "chunk_type": "image"
            }
            
            print(f"✓ Parsed image: {filename}")
            return [chunk]
            
        except Exception as e:
            print(f"Error parsing image {filename}: {e}")
            raise Exception(f"Failed to parse image: {str(e)}")


# Singleton instance
parser = ParserService()


if __name__ == "__main__":
    print("Testing Vextral Enhanced Parser...")
    
    test_text = "This is a test sentence. " * 300
    chunks = parser.chunk_text_hierarchical(test_text)
    print(f"✓ Hierarchical chunking: {len(chunks)} children linked to parents")
    print(f"✓ Supported formats: PDF, DOCX, TXT, CSV, MD, JSON, PNG, JPG, JPEG, WEBP")
    print("\n✓ ENHANCED PARSER READY")
